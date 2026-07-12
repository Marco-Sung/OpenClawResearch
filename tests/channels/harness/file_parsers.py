"""Parse-from-BYTES file extractors, for transport mode.
 
Your file_extractors.py builds an artifact from a spec and parses it in one
call -- fine for direct mode, but transport hands you the bytes of a file that
already crossed a channel, with no spec to look at. These mirror the same
strategy names so `attack["extract"]` lines up 1-to-1.
"""
from __future__ import annotations
 
import io
 
 
def _txt(raw: bytes) -> str:
    return raw.decode(errors="replace").strip()
 
 
def _pdf_naive(raw: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(raw))
    return "\n".join(p.extract_text() for p in reader.pages).strip()
 
 
def _pdf_color_filtered(raw: bytes) -> str:
    """Approximate color-aware extraction: track the current fill color via
    the content-stream operators and drop text drawn while fill is (near)
    white. This is the REAL version of the hardening your spec-shortcut faked
    -- and it can't cheat, because over transport there is no spec to fall
    back to. Approximate + UNTESTED on your PDFs; verify on the VM."""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(raw))
    out = []
    for page in reader.pages:
        state = {"white": False}
        kept: list[str] = []
 
        def before(op, args, cm, tm, _s=state):
            name = op.decode() if isinstance(op, (bytes, bytearray)) else str(op)
            try:
                if name == "g" and len(args) == 1:
                    _s["white"] = float(args[0]) >= 0.95
                elif name == "rg" and len(args) == 3:
                    _s["white"] = all(float(a) >= 0.95 for a in args)
                elif name == "k" and len(args) == 4:          # CMYK white = all 0
                    _s["white"] = all(float(a) <= 0.05 for a in args)
            except (TypeError, ValueError):
                pass
 
        def text(t, cm, tm, font, size, _s=state, _k=kept):
            if not _s["white"]:
                _k.append(t)
 
        page.extract_text(visitor_operand_before=before, visitor_text=text)
        out.append("".join(kept))
    return "\n".join(out).strip()
 
 
def _docx_naive(raw: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(raw))
    return "\n".join(p.text for p in doc.paragraphs).strip()
 
 
def _docx_visible_only(raw: bytes) -> str:
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(io.BytesIO(raw))
    out = []
    for para in doc.paragraphs:
        visible = []
        for run in para.runs:
            rPr = run._element.find(qn("w:rPr"))
            hidden = rPr is not None and rPr.find(qn("w:vanish")) is not None
            if not hidden:
                visible.append(run.text)
        out.append("".join(visible))
    return "\n".join(out).strip()
 
 
FILE_PARSERS = {
    "txt": _txt,
    "pdf_naive": _pdf_naive,
    "pdf_color_filtered": _pdf_color_filtered,
    "docx_naive": _docx_naive,
    "docx_visible_only": _docx_visible_only,
}
 
 
def parse_file_bytes(raw: bytes, strategy: str) -> str:
    if strategy not in FILE_PARSERS:
        raise KeyError(f"Unknown file strategy '{strategy}'. "
                       f"Options: {list(FILE_PARSERS)}")
    return FILE_PARSERS[strategy](raw)
 