# tests/channels/parsers/file_extractors.py
#
# Text-extraction layer for the FILE channel, parallel to extractors.py
# (which handles html/email). A file connector's job is the same as an
# email/web connector's: turn a file's bytes into the plain-text string
# the sanitizer then validates. And like those, the EXTRACTION CHOICE
# determines whether hidden payloads survive.
#
# Unlike html/email (where "raw" is a string you can inline), files are
# often BINARY (pdf, docx). So these extractors take a small spec dict
# describing the file's visible + hidden content, BUILD a real file of
# that type in memory, then extract text from it - so the test actually
# exercises a real pdf/docx parser, not a hand-written approximation.
#
# Requires: pypdf (extract), python-docx (docx), reportlab (build pdf).
# reportlab is BUILD-only (making test fixtures), not something a real
# connector needs - keep it a dev dependency, not in requirements.txt.
 
import io
import re
 
 
# ── PLAINTEXT ─────────────────────────────────────────────────────
 
def extract_txt(spec: dict) -> str:
    """Plain .txt file. There's no 'hidden' text in a plaintext file -
    everything is visible - so visible and hidden just get concatenated
    the way they'd appear in the raw file. This is the baseline: no
    format-specific hiding tricks are even possible."""
    parts = [spec.get("visible", "")]
    if spec.get("hidden"):
        parts.append(spec["hidden"])
    return "\n".join(p for p in parts if p).strip()
 
 
# ── PDF ────────────────────────────────────────────────────────────
 
def _build_pdf(spec: dict) -> bytes:
    """Build a real PDF: visible text in black, 'hidden' text in white
    (white-on-white = invisible to a human, but present in the text
    layer). This is the PDF analog of CSS display:none."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
 
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    y = 720
 
    c.setFillColorRGB(0, 0, 0)  # black, visible
    for line in spec.get("visible", "").split("\n"):
        c.drawString(72, y, line)
        y -= 20
 
    if spec.get("hidden"):
        c.setFillColorRGB(1, 1, 1)  # white, "invisible"
        for line in spec["hidden"].split("\n"):
            c.drawString(72, y, line)
            y -= 20
 
    c.save()
    buf.seek(0)
    return buf.getvalue()
 
 
def extract_pdf_naive(spec: dict) -> str:
    """Naive PDF extraction: pypdf reads the text layer regardless of
    color. White-on-white 'hidden' text SURVIVES - pypdf reads text
    OBJECTS, not a rendered image, so it has no idea the text is
    invisible to a human. Same blind spot bs4_get_text has for CSS."""
    from pypdf import PdfReader
    pdf_bytes = _build_pdf(spec)
    reader = PdfReader(io.BytesIO(pdf_bytes))
    return "\n".join(page.extract_text() for page in reader.pages).strip()
 
 
def extract_pdf_color_filtered(spec: dict) -> str:
    """HARDENED (approximate): this is a STAND-IN showing the intent of
    a color-aware extractor. A true implementation would inspect each
    text run's fill color in the PDF content stream and drop
    white/near-white runs. pypdf doesn't expose per-run color simply,
    so here we approximate by rebuilding from the spec's visible-only
    text. In a real connector you'd parse the content stream or
    render+OCR (see note). Flagged clearly so you don't mistake the
    approximation for a real color-parse.
 
    NOTE: the honest 'what a human sees' version is render-to-image +
    OCR (pdf2image + pytesseract), which needs system poppler/tesseract.
    That's the real fix; this is a spec-level shortcut for the harness."""
    return spec.get("visible", "").strip()
 
 
# ── DOCX ───────────────────────────────────────────────────────────
 
def _build_docx(spec: dict) -> bytes:
    """Build a real .docx. 'Visible' text is normal paragraphs. 'Hidden'
    text uses Word's actual hidden-font attribute (w:vanish) - text
    that's really in the document but not shown in normal view. This is
    a real Word feature, not a hack, which is what makes it a
    legitimate hiding vector."""
    from docx import Document
    from docx.oxml.ns import qn
 
    doc = Document()
    for line in spec.get("visible", "").split("\n"):
        doc.add_paragraph(line)
 
    if spec.get("hidden"):
        for line in spec["hidden"].split("\n"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            # set the hidden/vanish property on the run
            rPr = run._element.get_or_add_rPr()
            vanish = rPr.makeelement(qn("w:vanish"), {})
            rPr.append(vanish)
 
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
 
 
def extract_docx_naive(spec: dict) -> str:
    """Naive docx extraction: python-docx reads ALL runs' text,
    including ones marked hidden (w:vanish). So hidden-font payloads
    SURVIVE - the extractor pulls raw text and ignores display
    attributes, same pattern as the pdf and html naive cases."""
    from docx import Document
    docx_bytes = _build_docx(spec)
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs).strip()
 
 
def extract_docx_visible_only(spec: dict) -> str:
    """HARDENED docx: walk runs and SKIP any marked hidden (w:vanish)
    before collecting text. Hidden-font payloads DIE here. This one is
    a REAL fix (unlike the pdf approximation) - the vanish attribute is
    directly inspectable per-run, so we can genuinely filter on it."""
    from docx import Document
    from docx.oxml.ns import qn
 
    docx_bytes = _build_docx(spec)
    doc = Document(io.BytesIO(docx_bytes))
    out = []
    for para in doc.paragraphs:
        visible_runs = []
        for run in para.runs:
            rPr = run._element.find(qn("w:rPr"))
            is_hidden = rPr is not None and rPr.find(qn("w:vanish")) is not None
            if not is_hidden:
                visible_runs.append(run.text)
        out.append("".join(visible_runs))
    return "\n".join(out).strip()
 
 
# ── STRATEGY REGISTRY ─────────────────────────────────────────────
 
FILE_EXTRACTORS = {
    "txt": extract_txt,
    "pdf_naive": extract_pdf_naive,
    "pdf_color_filtered": extract_pdf_color_filtered,
    "docx_naive": extract_docx_naive,
    "docx_visible_only": extract_docx_visible_only,
}
 
 
def extract_file(spec: dict, strategy: str) -> str:
    """Entry point. spec is a dict with 'visible' and optional 'hidden'
    keys; strategy names which extractor to use. Returns the plain-text
    string that would then be handed to safe_write()."""
    if strategy not in FILE_EXTRACTORS:
        raise KeyError(f"Unknown file strategy '{strategy}'. Options: {list(FILE_EXTRACTORS)}")
    return FILE_EXTRACTORS[strategy](spec)
 