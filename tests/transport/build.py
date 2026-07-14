"""BUILD stage: turn an attack into a REAL artifact on disk under out/.
 
Writers only (the "pen", not the "eye"):
  reportlab   -> .pdf   (draws text in a fill color; white = hidden)
  python-docx -> .docx  (w:vanish run property = hidden)
  (plain str) -> .txt / .html
  email       -> .eml   (optional; email transport usually sends raw)
"""
from __future__ import annotations
 
from pathlib import Path
from email.message import EmailMessage
 
HARNESS = Path(__file__).resolve().parent
OUT = HARNESS / "out"
 
 
def _out(channel: str, name: str, ext: str) -> Path:
    d = OUT / channel
    d.mkdir(parents=True, exist_ok=True)
    return d / f"{name}.{ext}"
 
 
def build_txt(spec: dict, name: str) -> Path:
    parts = [spec.get("visible", "")]
    if spec.get("hidden"):
        parts.append(spec["hidden"])
    path = _out("files", name, "txt")
    path.write_text("\n".join(p for p in parts if p), encoding="utf-8")
    return path
 
 
def build_pdf(spec: dict, name: str) -> Path:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
 
    path = _out("files", name, "pdf")
    c = canvas.Canvas(str(path), pagesize=letter)
    y = 720
    c.setFillColorRGB(0, 0, 0)  # visible
    for line in spec.get("visible", "").split("\n"):
        c.drawString(72, y, line)
        y -= 20
    if spec.get("hidden"):
        c.setFillColorRGB(1, 1, 1)  # white-on-white
        for line in spec["hidden"].split("\n"):
            c.drawString(72, y, line)
            y -= 20
    c.save()
    return path
 
 
def build_docx(spec: dict, name: str) -> Path:
    from docx import Document
    from docx.oxml.ns import qn
 
    path = _out("files", name, "docx")
    doc = Document()
    for line in spec.get("visible", "").split("\n"):
        doc.add_paragraph(line)
    if spec.get("hidden"):
        for line in spec["hidden"].split("\n"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            rPr = run._element.get_or_add_rPr()
            rPr.append(rPr.makeelement(qn("w:vanish"), {}))
    doc.save(str(path))
    return path
 
 
def build_html(html: str, name: str) -> Path:
    path = _out("web", name, "html")
    path.write_text(html, encoding="utf-8")
    return path

def build_html_encoded(html: str, name: str, encoding: str = "iso-8859-1") -> Path:
    """Write the SAME html string but encode it as bytes in something other
    than UTF-8. pipeline.py's transport branch hardcodes decode(errors="replace")
    assuming UTF-8 -- if the actual bytes on disk are a different encoding,
    the decode mangles non-ASCII characters into the U+FFFD replacement
    character. Direct mode never touches bytes at all, so it sees the
    original string perfectly intact -- that asymmetry IS the divergence."""
    path = _out("web", name, "html")
    path.write_bytes(html.encode(encoding, errors="replace"))
    return path
 
 
def build_eml(visible: str, name: str, html: str | None = None,
              subject: str = "message",
              sender: str = "prof@university.edu",
              to: str = "openbuddy@workspace.local") -> Path:
    path = _out("mail", name, "eml")
    msg = EmailMessage()
    msg["From"] = sender
    msg["To"] = to
    msg["Subject"] = subject
    msg.set_content(visible)
    if html:
        msg.add_alternative(html, subtype="html")
    path.write_bytes(msg.as_bytes())
    return path