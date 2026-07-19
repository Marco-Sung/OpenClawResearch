"""Parse-from-BYTES file extractors, for transport mode.
 
Your extraction/files.py builds an artifact from a spec and parses it in one
call -- fine for direct mode, but transport hands you the bytes of a file that
already crossed a channel, with no spec to look at. These mirror the same
strategy names so `attack["extract"]` lines up 1-to-1.
"""
from __future__ import annotations
 
import io
 
 
def _txt(raw: bytes) -> str:
    return raw.decode(errors="replace").strip()
 
 
def _pdf_naive(raw: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(raw))
    return "\n".join(p.extract_text() for p in reader.pages).strip()
 
 
def _pdf_color_filtered(raw: bytes) -> str:
    """Real color-aware extraction -- now shared with the direct path via
    extraction/pdf_color.py (step 4c), so direct and transport modes give
    identical, genuine hardening instead of disagreeing."""
    from tests.extraction.pdf_color import filter_white_text
    return filter_white_text(raw)
 
 
def _docx_naive(raw: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(raw))
    return "\n".join(p.text for p in doc.paragraphs).strip()
 
 
def _docx_visible_only(raw: bytes) -> str:
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(io.BytesIO(raw))
    out = []
    for para in doc.paragraphs:
        visible = []
        for run in para.runs:
            rPr = run._element.find(qn("w:rPr"))
            hidden = rPr is not None and rPr.find(qn("w:vanish")) is not None
            if not hidden:
                visible.append(run.text)
        out.append("".join(visible))
    return "\n".join(out).strip()
 
 
FILE_PARSERS = {
    "txt": _txt,
    "pdf_naive": _pdf_naive,
    "pdf_color_filtered": _pdf_color_filtered,
    "docx_naive": _docx_naive,
    "docx_visible_only": _docx_visible_only,
}
 
 
def parse_file_bytes(raw: bytes, strategy: str) -> str:
    if strategy not in FILE_PARSERS:
        raise KeyError(f"Unknown file strategy '{strategy}'. "
                       f"Options: {list(FILE_PARSERS)}")
    return FILE_PARSERS[strategy](raw)
 